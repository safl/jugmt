"""
A representation of Figure data extracted from a ``.docx`` document

This is intended to be used with NVMe Specfication Documents, however, if your
assumptions match those below, then it will work with your documents as well.

Assumptions
===========

* The first cell in the first row contains a Figure caption on the form

  - 'Figure <figure_number>: <description>'

* The list of figures contains figure captions on the form

  - 'Figure <figure_number>: <description> <page_nr>'
"""

import re
from pathlib import Path
from typing import ClassVar, Dict, List, Optional

import docx
from jinja2 import Environment, PackageLoader, select_autoescape
from pydantic import BaseModel, Field, ValidationError

import jugmt


def pascal_to_snake(name):
    """Convert a PascalCase to snake_case"""
    return re.sub(r"(?<!^)(?=[A-Z])", "_", name).lower()


def snake_to_pascal(name):
    """Convert snake_case to PascalCase"""

    return "".join(word.capitalize() for word in name.split("_"))


class Cell(BaseModel):
    key: str = "cell"
    text: str = Field(default_factory=str)
    tables: List["Table"] = Field(default_factory=list)


class Row(BaseModel):
    key: str = "row"
    cells: List[Cell] = Field(default_factory=list)


class Table(BaseModel):
    key: str = "table"
    rows: List[Row] = Field(default_factory=list)

    @classmethod
    def from_docx_table(cls, docx_table: docx.table.Table):
        table = cls()

        for docx_row in docx_table.rows:
            row = Row()
            table.rows.append(row)

            for docx_cell in docx_row.cells:
                cell = Cell(
                    text=str(docx_cell.text),
                    tables=[
                        Table.from_docx_table(nested_table)
                        for nested_table in docx_cell.tables
                    ],
                )
                row.cells.append(cell)

        return table


class Figure(BaseModel):
    """
    A figure as captioned in the NVMe Specification Documents.

    Ensuring that the 'figure' being "enriched" is always available. Providing a
    factory-method, constructing the class using 'figure' information, and enriched data
    from 'match.groupdict()'
    """

    REGEX_FIGURE_CAPTION: ClassVar[str] = (
        r"^(?P<caption>Figure\s+(?P<figure_nr>\d+)\s*:"
        r"\s*(?P<description>.*?))(?P<page_nr>\d+)?$"
    )

    key: str  # Utilized fo de-serialization and document location
    figure_nr: int  # Figure as numbered in the specification document
    caption: str  # The entire figure title
    description: str  # The part of figure title without the "Fig X:" prefix

    page_nr: Optional[int] = None
    table: Optional[Table] = None

    @classmethod
    def from_figure_caption(cls, text: str):
        match = re.match(Figure.REGEX_FIGURE_CAPTION, text)
        if not match:
            return None

        args = {}
        args["key"] = pascal_to_snake(cls.__name__)
        args["figure_nr"] = int(match.group("figure_nr"))
        args["caption"] = match.group("caption").strip()
        args["description"] = match.group("description").strip()
        args["page_nr"] = (
            int(match.group("page_nr")) if match.group("page_nr") else None
        )
        args["table"] = None

        return cls(**args)


class Meta(BaseModel):
    key: str = "meta"
    version: str = jugmt.__version__
    stem: str = Field(default_factory=str)


class Document(BaseModel):
    """
    Wrapper of the docx.Document with additional path meta-data and figures with tabular
    data converted to a JSON serializable format
    """

    TEMPLATE_HTML: ClassVar[str] = "document.figures.html.jinja2"

    meta: Meta = Field(default_factory=Meta)
    figures: List[Figure] = Field(default_factory=list)

    @classmethod
    def from_docx(cls, path: Path):
        """
        Populate self.figures with figures, and their associated tabular data when
        available, found in self.docx
        """

        docx_document = docx.Document(path)
        figures = {}

        errors: Dict[str, List[tuple]] = {
            "captions": [],
            "tof_entries": [],
        }

        # Add tabular figures -- page_nr unavailable
        for table_nr, docx_table in enumerate(docx_document.tables, 1):
            caption = str(docx_table.rows[0].cells[0].text).strip()

            figure = Figure.from_figure_caption(caption)
            if not figure:
                errors["captions"].append(
                    (table_nr, caption, "Does not match figure caption assumptions")
                )
                continue

            if figure.figure_nr in figures:
                errors["captions"].append((table_nr, caption, "Duplicate"))
                continue

            figure.table = Table.from_docx_table(docx_table)
            figures[figure.figure_nr] = figure

        # Update tabular figures with page_nr
        # Add non-fabular figures
        # Check table-of-figure description validity
        prev = cur = None
        for paragraph in docx_document.paragraphs:
            cur = paragraph.style.name

            # We exit early to avoid scanning the entire document, since we know that
            # once we are looking at a "table of figures" paragraph, then once we see
            # one that is not, then no more will arrive
            if prev == "table of figures" and cur != "table of figures":
                break
            prev = cur
            if paragraph.style.name != "table of figures":
                continue

            # Check whether the paragraph is a reference to a figure
            caption = paragraph.text.strip()
            figure = Figure.from_figure_caption(caption)
            if not figure:
                errors["tof_entries"].append(
                    (caption, "Does not match figure assumptions")
                )
                continue

            if not figure.page_nr:
                errors["tof_entries"].append((caption, "Is missing <page_nr>"))
                continue

            existing = figures.get(figure.figure_nr, None)
            if existing:
                existing.page_nr = figure.page_nr
                if figure.description not in existing.description:
                    errors["tof_entries"].append(
                        (
                            caption,
                            f"({existing.description}) != {figure.description}",
                        )
                    )
            else:
                figures[figure.figure_nr] = figure

        return cls(meta=Meta(stem=path.stem), figures=list(figures.values())), errors

    def to_html(self) -> str:
        """Returns the document as a HTML-formatted string"""

        env = Environment(
            loader=PackageLoader("jugmt", "templates"),
            autoescape=select_autoescape(["html", "xml"]),
        )
        template = env.get_template(Document.TEMPLATE_HTML)

        return template.render(document=self)

    def to_json(self) -> str:
        """Returns the document as a JSON-formatted string"""

        return self.model_dump_json(indent=4)

    def check(self):
        try:
            self.validate(self.dict())
        except ValidationError as e:
            print(e)
            return False

        return True
