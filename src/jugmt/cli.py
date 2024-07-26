"""
Command-Line Interface
======================

Produces .json and .html when given NVMe Specicification document in .docx format
"""

from argparse import ArgumentParser, Namespace
from pathlib import Path
from typing import Any, Dict, List, Tuple

import docx
from senfd.documents import DocumentMeta, FigureDocument
from senfd.figures import Figure
from senfd.tables import Cell, Row, Table

import jugmt


def docx_table_to_table(docx_table: docx.table.Table) -> Table:
    table = Table()

    for docx_row in docx_table.rows:
        row = Row()
        table.rows.append(row)

        for docx_cell in docx_row.cells:
            cell = Cell(
                text=str(docx_cell.text),
                tables=[
                    docx_table_to_table(nested_table)
                    for nested_table in docx_cell.tables
                ],
            )
            row.cells.append(cell)

    return table


def docx_to_figure_document(path: Path) -> Tuple[FigureDocument, Dict[str, Any]]:
    """
    Returns (figure_document, errors)
    """

    figures = {}
    errors: Dict[str, List[tuple]] = {
        "captions": [],
        "tof_entries": [],
    }

    docx_document = docx.Document(path)

    # Add tabular figures -- page_nr unavailable
    for table_nr, docx_table in enumerate(docx_document.tables, 1):
        caption = str(docx_table.rows[0].cells[0].text).strip()

        figure = Figure.from_regex(Figure.REGEX_TABLE_ROW, caption)
        if not figure:
            errors["captions"].append(
                (table_nr, caption, "Does not match figure caption assumptions")
            )
            continue

        if figure.figure_nr in figures:
            errors["captions"].append((table_nr, caption, "Duplicate"))
            continue

        figure.table = docx_table_to_table(docx_table)
        figures[figure.figure_nr] = figure

    # Update tabular figures with page_nr
    # Add non-fabular figures
    # Check table-of-figure description validity
    prev = cur = None
    for paragraph in docx_document.paragraphs:
        cur = paragraph.style.name

        # We exit early to avoid scanning the entire document, since we know that
        # once we are looking at a "table of figures" paragraph, then once we see
        # one that is not, then no more will arrive
        if prev == "table of figures" and cur != "table of figures":
            break
        prev = cur
        if paragraph.style.name != "table of figures":
            continue

        # Check whether the paragraph is a reference to a figure
        caption = paragraph.text.strip()
        figure = Figure.from_regex(Figure.REGEX_TABLE_OF_FIGURES, caption)
        if not figure:
            errors["tof_entries"].append((caption, "Does not match figure assumptions"))
            continue

        if not figure.page_nr:
            errors["tof_entries"].append((caption, "Is missing <page_nr>"))
            continue

        existing = figures.get(figure.figure_nr, None)
        if existing:
            existing.page_nr = figure.page_nr
            if figure.description not in existing.description:
                errors["tof_entries"].append(
                    (
                        caption,
                        f"({existing.description}) != {figure.description}",
                    )
                )
        else:
            figures[figure.figure_nr] = figure

    return (
        FigureDocument(
            meta=DocumentMeta(stem=path.stem), figures=list(figures.values())
        ),
        errors,
    )


def parse_args() -> Namespace:
    """Return command-line arguments"""

    parser = ArgumentParser(
        description="Extract table information from .docx documents"
    )
    parser.add_argument(
        "document", nargs="*", type=Path, help="path to one or more .docx document(s)"
    )
    parser.add_argument(
        "--output",
        type=Path,
        help="directory where the output will be saved",
        default=Path("output"),
    )
    parser.add_argument(
        "--version",
        action="store_true",
        help="print the version and exit",
    )

    args = parser.parse_args()
    if not args.document and not args.version:
        parser.error("the following arguments are required: document")

    return args


def main() -> int:
    """Command-line entrypoint"""

    args = parse_args()
    args.output.mkdir(parents=True, exist_ok=True)

    if args.version:
        print(jugmt.__version__)
        return 0

    for path in args.document:
        document, errors = docx_to_figure_document(path)
        document.to_html_file(args.output)
        document.to_json_file(args.output)

    return 0
